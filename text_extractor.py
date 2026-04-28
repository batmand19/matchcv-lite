import io
import unicodedata


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts text from PDF, DOCX, or TXT files.
    Returns a string with the extracted text.
    Raises RuntimeError with a human-readable message on any failure.
    """
    if not file_bytes:
        raise RuntimeError("El archivo está vacío. Sube un archivo con contenido.")

    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    try:
        if ext == "pdf":
            return _extract_from_pdf(file_bytes)
        elif ext == "docx":
            return _extract_from_docx(file_bytes)
        elif ext == "txt":
            return _extract_from_txt(file_bytes)
        else:
            raise ValueError(f"Formato de archivo no soportado: .{ext}. Usa PDF, DOCX o TXT.")
    except RuntimeError:
        raise
    except ValueError as e:
        raise RuntimeError(str(e))
    except Exception as e:
        raise RuntimeError(f"Error al leer el archivo '{filename}': {type(e).__name__} — {str(e)}")


def _extract_from_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2
    except ImportError:
        raise RuntimeError("La librería PyPDF2 no está instalada. Ejecuta: pip install PyPDF2")

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    except PyPDF2.errors.PdfReadError as e:
        raise RuntimeError(f"El PDF está dañado o protegido y no puede leerse: {e}")
    except Exception as e:
        raise RuntimeError(f"No se pudo abrir el PDF: {type(e).__name__} — {e}")

    if len(reader.pages) == 0:
        raise RuntimeError("El PDF no contiene páginas.")

    text_parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())
        except Exception:
            # Skip unreadable pages silently; we'll catch empty result below
            pass

    result = "\n".join(text_parts).strip()
    if not result:
        raise RuntimeError(
            "No se pudo extraer texto del PDF. "
            "Es posible que sea un PDF escaneado (imagen). "
            "Intenta copiando el texto manualmente en la pestaña 'Pegar texto del CV'."
        )
    return result


def _extract_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        raise RuntimeError("La librería python-docx no está instalada. Ejecuta: pip install python-docx")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except PackageNotFoundError:
        raise RuntimeError(
            "El archivo .docx está dañado o no es un Word válido. "
            "Guárdalo de nuevo desde Word o Google Docs y vuelve a subirlo."
        )
    except Exception as e:
        raise RuntimeError(f"No se pudo abrir el archivo DOCX: {type(e).__name__} — {e}")

    text_parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_parts.append(t)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    text_parts.append(t)

    result = "\n".join(text_parts).strip()
    if not result:
        raise RuntimeError(
            "El archivo DOCX no tiene texto extraíble. "
            "Puede estar vacío o contener solo imágenes."
        )
    return result


def _extract_from_txt(file_bytes: bytes) -> str:
    if len(file_bytes) == 0:
        raise RuntimeError("El archivo TXT está vacío.")

    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding).strip()
            if text:
                return text
        except (UnicodeDecodeError, LookupError):
            continue

    raise RuntimeError(
        "No se pudo decodificar el archivo TXT. "
        "Asegúrate de que esté guardado en UTF-8 o Latin-1."
    )
